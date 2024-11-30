import io
import os
import json
import docx

from fastapi import APIRouter
from fastapi.responses import StreamingResponse

from pymongo import MongoClient
from typing import Dict
from pydantic import BaseModel

from common.llm.models import SamplingParams
from common.llm.utils import change_current_lora_adapter, get_completion

from common.db.utils import upload_document
from common.db.models import *


class InputModel(BaseModel):
    """
    A model to represent input to **/generate** method.

    Attributes
    ----------
    type : str
        Type of the document: "Акты", "Жалобы", etc.

    fields : Dict[str, str]
        Dynamic fields of the specific document type.

    params : SamplingParams
        Params to LLM generation: max_new_tokens, temperature, etc. 
    """
    
    type: str
    fields: Dict[str, str]
    params: SamplingParams


def text_to_docx_bytes(text: str) -> bytes:
    """
    Return .docx file in bytes.

    Parameters
    ----------
    text : str
        Content of document.

    Returns
    -------
    bytes
        Bytes of document.
    """

    doc = docx.Document()
    
    doc.add_paragraph(text)
    
    byte_stream = io.BytesIO()
    
    doc.save(byte_stream)
    
    byte_stream.seek(0)
    return byte_stream.getvalue()


router = APIRouter()


@router.post("/generate")
async def generate(input: InputModel):
    # change current lora adapter if need:
    change_current_lora_adapter(adapter_id=int(os.environ.get("GENERATE_LORA_ADAPTER_ID")))
    # generate document based on fields:
    document_as_text = get_completion(prompt=json.dumps(input.fields, ensure_ascii=False), params=input.params.model_dump())
    
    # transform document in text format to byte
    document_as_byte = text_to_docx_bytes(document_as_text)
    
    # upload document to mongodb
    upload_document(
        Document(
            file=document_as_byte, 
            info=DocumentInfo(
                header=input.type, description=", ".join(list(input.fields.values()))
            ),
            type=DocumentType(
                label=input.type, score=1.0
            ),
            entities=[Entity(label=key, value=value) for key, value in input.fields.items()]
        )
    )

    # return document from FastAPI
    file_like_document = io.BytesIO(document_as_byte)
    file_like_document.seek(0)

    return StreamingResponse(file_like_document, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=document.docx"})

