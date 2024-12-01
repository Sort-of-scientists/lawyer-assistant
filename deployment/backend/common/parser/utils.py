import io
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from typing import Any
from tika import parser

TIKA_ENDPOINT = "http://tika:9998/tika"


def text_to_docx_bytes(text: str, md: bool = True) -> bytes:
    """
    Converts a markdown-formatted string into a .docx document and returns it as bytes.

    The function creates a Document object, adds the parsed content from the markdown text,
    and saves the document into a byte stream.

    Parameters
    ----------
    text : str
        Content of document.
    md : bool, optional
        If set to 'True' (default), the text is interpreted as Markdown and parsed accordingly.
        If 'False', the text is added as a single paragraph without any special formatting.

    Returns
    -------
    bytes
        Bytes of document.
    """

    doc = Document()
    
    if md:
        parse_markdown(doc, text)
    else:
        doc.add_paragraph(text)
    
    byte_stream = io.BytesIO()
    
    doc.save(byte_stream)
    
    byte_stream.seek(0)
    return byte_stream.getvalue()


def add_heading_with_style(doc: Document, text: str, level: int) -> Paragraph:
    """
    Adds a heading to the document with specified level, aligns it to the center,
    and sets its font color to black.

    Parameters
    ----------
    doc : Document 
        The document object where the heading will be added.
    text : str 
        The text content of the heading.
    level : int 
        The level of the heading (e.g., 1 for H1, 2 for H2, etc.).

    Returns
    ----------
    heading : Paragraph 
        A paragraph object representing the newly added heading.
    """
    
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    return heading


def add_formatted_paragraph(doc: Document, text: str) -> Paragraph:
    """
    Adds a formatted paragraph to the given Document object based on the input text.

    The input text can contain markdown-like formatting such as:
    - Bold text: denoted by double asterisks (e.g., **text**)
    - Italic text: denoted by single asterisks (e.g., *text*)

    Parameters
    ----------
    doc : Document 
        The Document object to which the paragraph will be added.
    text : str 
        The text containing markdown-like formatting.

    Returns
    ----------
    prg : Paragraph
        The paragraph object that was added to the Document.
    """
    
    prg = doc.add_paragraph()
    
    i = 0
    while i < len(text):
        if i + 1 < len(text) and text[i] == '*' and text[i + 1] == '*': 
            j = i + 2
            while j < len(text) and not (text[j] == '*' and text[j + 1] == '*'):
                j += 1
            prg.add_run(text[i + 2:j]).bold = True
            i = j + 2
        elif i < len(text) and text[i] == '*':
            j = i + 1
            while j < len(text) and text[j] != '*':
                j += 1
            prg.add_run(text[i + 1:j]).italic = True
            i = j + 1
        else:
            prg.add_run(text[i])
            i += 1
            
    prg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    return prg


def parse_markdown(doc: Document, markdown_text: str):    
    """
    Parses a markdown-formatted string and adds the content to the given Document object.

    Parameters
    ----------
    doc : Document
        The Document object to which the parsed content will be added.
    markdown_text : str 
        The markdown-formatted text to be parsed and added to the Document.
    """
    
    lines = markdown_text.split('\n')

    for line in lines:
        
        if line.startswith('#'):
            level = line.count('#')
            text = line[level + 1:].strip() 
            add_heading_with_style(doc, text, level)
        else:
            add_formatted_paragraph(doc, line)


def parse_document(document_as_bytes: Any) -> str:
    return parser.from_buffer(document_as_bytes, TIKA_ENDPOINT)["content"]