from fastapi import FastAPI
from fastapi.responses import FileResponse
from docx import Document
import os
import uuid
import mistune

app = FastAPI()


@app.post("/create-document")
async def convert_markdown_to_docx(markdown_text: str):
    # Generate a unique filename for the output .docx file
    output_filename = f"{uuid.uuid4()}.docx"

    # Parse the Markdown content
    markdown_parser = mistune.create_markdown()
    html_content = markdown_parser(markdown_text)

    # Create a new .docx document
    doc = Document()

    # Simple markdown-to-docx conversion (basic handling of paragraphs and headings)
    from html.parser import HTMLParser

    class MyHTMLParser(HTMLParser):
        def __init__(self, document):
            super().__init__()
            self.document = document

        def handle_starttag(self, tag, attrs):
            if tag == "h1":
                self.document.add_heading(level=1)
            elif tag == "h2":
                self.document.add_heading(level=2)
            elif tag == "h3":
                self.document.add_heading(level=3)

        def handle_data(self, data):
            self.document.add_paragraph(data)

    # Create a custom HTML parser to convert parsed Markdown (HTML) into .docx
    parser = MyHTMLParser(doc)
    parser.feed(html_content)

    # Save the document to a .docx file
    doc.save(output_filename)

    # Return the generated .docx file as a response
    return FileResponse(output_filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=output_filename)
